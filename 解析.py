from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_abstract_document(file_name, abstract_content, keywords):
    # 创建文档
    doc = Document()

    # 添加标题
    title = doc.add_paragraph('摘要', style='Heading 1')
    title.style.font.bold = False  # 不加粗
    title.style.font.size = Pt(16)  # 黑体三号字
    title.style.font.name = '黑体'
    title.style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_before = Pt(2 * 12)  # 段前2行
    title.paragraph_format.space_after = Pt(2 * 12)  # 段后2行

    # 添加摘要内容
    abstract_paragraphs = abstract_content.split('\n')  # 分隔成多个段落
    for paragraph_text in abstract_paragraphs:
        abstract = doc.add_paragraph(paragraph_text, style='BodyText')
        abstract.style.font.size = Pt(12)
        abstract.paragraph_format.line_spacing = Pt(20)  # 固定行距为20磅
        abstract.style.font.name = '宋体'
        # 设置首行缩进两行
        abstract.paragraph_format.first_line_indent = Pt(2 * 12)  # 每行12磅，两行就是24磅

    # 添加空行
    doc.add_paragraph()

    # 添加关键词内容
    keyword_paragraph = doc.add_paragraph('关键词：', style='BodyText')
    keyword_paragraph.style.font.bold = False  # 不加粗
    keyword_paragraph.style.font.size = Pt(12)
    keyword_paragraph.style.font.name = '黑体'

    # 添加关键词内容
    keyword_run = keyword_paragraph.add_run(keywords)
    keyword_run.style.font.size = Pt(12)
    keyword_run.style.font.name = '宋体'

    # 设置关键词段落的行距为20磅
    keyword_paragraph.paragraph_format.line_spacing = Pt(20)  # 固定行距为20磅

    # 保存文档
    doc.save(file_name)

# 示例调用
create_abstract_document('abstract1.docx', '这是摘要内容。\n这是摘要的第二段内容。', '关键词1;关键词2;关键词3')
